"""DOCX document parser using python-docx."""
from pathlib import Path
from typing import Union, List
from uuid import UUID
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.parsing.base import BaseParser
from app.services.parsing.models import ParsedDocument, ParsedSection
from app.services.parsing.normalizer import TextNormalizer
from app.core.logging import get_logger

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents using python-docx."""
    
    def __init__(self):
        """Initialize DOCX parser."""
        self.normalizer = TextNormalizer()
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this parser supports DOCX format."""
        return file_extension.lower() in ['docx']
    
    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a heading."""
        return paragraph.style.name.startswith('Heading')
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """Get heading level (1-9)."""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            try:
                # Extract number from 'Heading 1', 'Heading 2', etc.
                level = int(style_name.split()[-1])
                return level
            except (ValueError, IndexError):
                return 1
        return 0
    
    def _table_to_markdown(self, table: Table) -> str:
        """
        Convert DOCX table to markdown format.
        
        Args:
            table: DOCX table object
            
        Returns:
            str: Markdown representation of table
        """
        if not table.rows:
            return ""
        
        lines = []
        
        # Process header row
        if table.rows:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            lines.append('| ' + ' | '.join(header_cells) + ' |')
            lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
        
        # Process data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(lines)
    
    async def parse(
        self,
        file_path: Union[str, Path],
        document_id: UUID,
    ) -> ParsedDocument:
        """
        Parse DOCX document.
        
        Features:
        - Detect and preserve headings
        - Preserve paragraph separation
        - Convert tables to markdown
        - Preserve bullet lists
        
        Args:
            file_path: Path to DOCX file
            document_id: Document UUID
            
        Returns:
            ParsedDocument: Parsed DOCX with sections
            
        Raises:
            Exception: If DOCX parsing fails
        """
        file_path = Path(file_path)
        
        logger.info(
            f"Parsing DOCX document",
            extra={
                "document_id": str(document_id),
                "file_path": str(file_path),
            }
        )
        
        try:
            # Open DOCX document
            doc = Document(file_path)
            
            sections: List[ParsedSection] = []
            current_section_title = None
            current_content_parts = []
            
            def save_current_section():
                """Save current accumulated section."""
                if current_content_parts:
                    content = '\n\n'.join(current_content_parts)
                    normalized_content = self.normalizer.normalize(content)
                    
                    if normalized_content:
                        section = ParsedSection(
                            section_title=current_section_title,
                            content=normalized_content,
                            metadata={
                                "paragraph_count": len(current_content_parts),
                            }
                        )
                        sections.append(section)
            
            # Process document elements
            for element in doc.element.body:
                # Check if element is a paragraph
                if element.tag.endswith('p'):
                    para = Paragraph(element, doc)
                    
                    # Skip empty paragraphs
                    if not para.text.strip():
                        continue
                    
                    # Check if it's a heading
                    if self._is_heading(para):
                        # Save previous section
                        save_current_section()
                        
                        # Start new section with heading
                        current_section_title = self.normalizer.normalize_section_title(para.text)
                        current_content_parts = []
                    else:
                        # Add paragraph to current section
                        current_content_parts.append(para.text)
                
                # Check if element is a table
                elif element.tag.endswith('tbl'):
                    table = Table(element, doc)
                    markdown_table = self._table_to_markdown(table)
                    
                    if markdown_table:
                        current_content_parts.append(markdown_table)
            
            # Save final section
            save_current_section()
            
            # If no sections were created (no headings), create one section with all content
            if not sections:
                all_text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                normalized_text = self.normalizer.normalize(all_text)
                
                if normalized_text:
                    section = ParsedSection(
                        section_title=None,
                        content=normalized_text,
                        metadata={
                            "paragraph_count": len(doc.paragraphs),
                        }
                    )
                    sections.append(section)
            
            # Create parsed document
            parsed_doc = ParsedDocument(
                document_id=document_id,
                sections=sections,
                metadata={
                    "parser": "docx",
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                }
            )
            
            # Log parsing statistics
            self._log_parsing_stats(
                document_id=document_id,
                sections_count=len(sections),
                total_chars=parsed_doc.get_total_content_length(),
            )
            
            return parsed_doc
            
        except Exception as e:
            logger.error(
                f"Failed to parse DOCX: {str(e)}",
                extra={
                    "document_id": str(document_id),
                    "file_path": str(file_path),
                },
                exc_info=True,
            )
            raise Exception(f"DOCX parsing failed: {str(e)}")
